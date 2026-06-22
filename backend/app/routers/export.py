"""导出：Excel（.xlsx）与 Word（.docx），按分类层级组织，列序与主表一致。"""
import io

from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from ..database import get_db
from ..security import require_admin
from ..models import Classification, Indicator, IndicatorStatus, User
from ..utils import classification_path

router = APIRouter(prefix="/export", tags=["导出"])

LEVELS = ["一级分类", "二级分类", "三级分类"]
HEADERS = ["来源标准/部分", *LEVELS, "标识符", "中文名称", "英文名称", "计量单位",
           "定义", "计算方法", "指标说明", "调查方法", "数据来源", "发布频率",
           "分层统计", "来源标签", "指标类型"]


def source_tags_str(ind):
    parts = []
    for t in (ind.source_tags or []):
        if t == "其他" and (ind.source_other or "").strip():
            parts.append(f"其他（{ind.source_other.strip()}）")
        else:
            parts.append(t)
    return "、".join(parts)


def _walk(db: Session, parent_id=None, path=None):
    """深度优先遍历分类树，产出 (node, [一级,二级,三级名称]) 序列。"""
    path = path or []
    nodes = (db.query(Classification).filter(Classification.parent_id == parent_id)
             .order_by(Classification.sort_order, Classification.id).all())
    for n in nodes:
        p = path + [n.name]
        yield n, p
        yield from _walk(db, n.id, p)


def _indicators(db: Session, class_id: int):
    return (db.query(Indicator)
            .filter(Indicator.classification_id == class_id, Indicator.status == IndicatorStatus.active)
            .order_by(Indicator.sort_order, Indicator.identifier).all())


@router.get("/excel", summary="导出 Excel")
def export_excel(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook(); ws = wb.active; ws.title = "卫生统计指标"
    thin = Side(style="thin", color="D0D0D0"); border = Border(thin, thin, thin, thin)
    ws.append(HEADERS)
    for c in range(1, len(HEADERS) + 1):
        cell = ws.cell(1, c)
        cell.font = Font(name="宋体", size=10, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E5F")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    for node, path in _walk(db):
        for ind in _indicators(db, node.id):
            levels = (path + ["", "", ""])[:3]
            ws.append([
                ind.source_standard.title if ind.source_standard else "",
                levels[0], levels[1], levels[2],
                ind.identifier, ind.name_cn, ind.name_en, ind.unit,
                ind.definition, ind.method, ind.description,
                ind.survey_method, ind.data_source, ind.frequency,
                ind.stratification, source_tags_str(ind), ind.indicator_type,
            ])
            r = ws.max_row
            for c in range(1, len(HEADERS) + 1):
                cc = ws.cell(r, c)
                cc.font = Font(name="宋体", size=10)
                cc.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                cc.border = border

    widths = [28, 12, 12, 12, 14, 22, 30, 8, 40, 32, 40, 10, 16, 10, 30, 24, 12]
    from openpyxl.utils import get_column_letter
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": "attachment; filename=health_indicators.xlsx"})


@router.get("/word", summary="导出 Word")
def export_word(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # 每个指标按此顺序，逐行输出（左列字段名，右列内容）
    FIELDS = [
        ("标识符", lambda i: i.identifier),
        ("中文名称", lambda i: i.name_cn),
        ("英文名称", lambda i: i.name_en),
        ("计量单位", lambda i: i.unit),
        ("定义", lambda i: i.definition),
        ("计算方法", lambda i: i.method),
        ("指标说明", lambda i: i.description),
        ("调查方法", lambda i: i.survey_method),
        ("数据来源", lambda i: i.data_source),
        ("发布频率", lambda i: i.frequency),
        ("分层统计", lambda i: i.stratification),
        ("来源标签", lambda i: source_tags_str(i)),
        ("指标类型", lambda i: i.indicator_type),
    ]

    def style_run(run, size=10.5, bold=False, color=None):
        """中文宋体、英文 Times New Roman。"""
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    def write_cell(cell, text, bold=False):
        cell.text = ""
        run = cell.paragraphs[0].add_run("" if text is None else str(text))
        style_run(run, bold=bold)

    doc = Document()
    # 页边距收窄，给内容列更多宽度
    sec = doc.sections[0]
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    # 文档默认样式：中文宋体、英文 Times New Roman
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    title = doc.add_heading(level=0)
    style_run(title.add_run("卫生统计指标（含元数据）"), size=18, bold=True, color=RGBColor(0x1F, 0x4E, 0x5F))

    def add_heading(text, level):
        h = doc.add_heading(level=level)
        style_run(h.add_run(text), size={1: 15, 2: 13, 3: 12, 4: 12}.get(level, 12),
                  bold=True, color=RGBColor(0x0F, 0x76, 0x6E))

    def emit(parent_id=None, depth=1):
        nodes = (db.query(Classification).filter(Classification.parent_id == parent_id)
                 .order_by(Classification.sort_order, Classification.id).all())
        for n in nodes:
            add_heading(n.name, min(depth, 4))
            for ind in _indicators(db, n.id):
                tbl = doc.add_table(rows=0, cols=2)
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                for label, getter in FIELDS:
                    cells = tbl.add_row().cells
                    write_cell(cells[0], label, bold=True)
                    write_cell(cells[1], getter(ind))
                    cells[0].width = Cm(2.2)
                    cells[1].width = Cm(14.8)
                doc.add_paragraph()  # 指标间留白
            emit(n.id, depth + 1)

    emit()
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=health_indicators.docx"})


# ---------------- 变更清单导出（新增 / 修改 / 删除） ----------------
def _changed_rows(db: Session, type: str):
    from .indicators import change_map, CHANGE_LABELS
    cmap = change_map(db)
    wanted = {"added", "modified", "deleted"} if type == "all" else {type}
    order = {"added": 0, "modified": 1, "deleted": 2}
    rows = [(ind, cmap.get(ind.id, "none")) for ind in db.query(Indicator).all()]
    rows = [(ind, ct) for ind, ct in rows if ct in wanted]
    rows.sort(key=lambda x: (order.get(x[1], 9), x[0].sort_order, x[0].identifier or ""))
    return rows, CHANGE_LABELS


@router.get("/changes/excel", summary="导出变更清单 Excel（管理员）")
def export_changes_excel(type: str = Query("all", pattern="^(all|added|modified|deleted)$"),
                         db: Session = Depends(get_db), _: User = Depends(require_admin)):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    rows, labels = _changed_rows(db, type)
    wb = Workbook(); ws = wb.active; ws.title = "变更清单"
    thin = Side(style="thin", color="D0D0D0"); border = Border(thin, thin, thin, thin)
    headers = ["变更类型", *HEADERS]
    ws.append(headers)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(1, c)
        cell.font = Font(name="宋体", size=10, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E5F")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    for ind, ct in rows:
        path = classification_path(db, ind.classification_id)
        levels = (path + ["", "", ""])[:3]
        ws.append([
            labels.get(ct, ct), ind.source_standard.title if ind.source_standard else "",
            levels[0], levels[1], levels[2],
            ind.identifier, ind.name_cn, ind.name_en, ind.unit, ind.definition, ind.method,
            ind.description, ind.survey_method, ind.data_source, ind.frequency,
            ind.stratification, source_tags_str(ind), ind.indicator_type,
        ])
        r = ws.max_row
        for c in range(1, len(headers) + 1):
            cc = ws.cell(r, c)
            cc.font = Font(name="宋体", size=10)
            cc.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cc.border = border
    for idx, w in enumerate([10, 28, 12, 12, 12, 14, 22, 30, 8, 40, 32, 40, 10, 16, 10, 30, 24, 12], start=1):
        ws.column_dimensions[chr(64 + idx) if idx <= 26 else "A" + chr(64 + idx - 26)].width = w
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": "attachment; filename=indicator_changes.xlsx"})


@router.get("/changes/word", summary="导出变更清单 Word（管理员）")
def export_changes_word(type: str = Query("all", pattern="^(all|added|modified|deleted)$"),
                        db: Session = Depends(get_db), _: User = Depends(require_admin)):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    rows, labels = _changed_rows(db, type)
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2)
    doc.add_heading("卫生统计指标变更清单", level=0)
    FIELDS = [("标识符", "identifier"), ("所属分类", "_path"), ("英文名称", "name_en"), ("计量单位", "unit"),
              ("定义", "definition"), ("计算方法", "method"), ("指标说明", "description"),
              ("调查方法", "survey_method"), ("数据来源", "data_source"), ("发布频率", "frequency"),
              ("分层统计", "stratification"), ("来源标签", "_tags"), ("指标类型", "indicator_type")]
    cur = None
    for ind, ct in rows:
        if ct != cur:
            doc.add_heading(f"{labels.get(ct, ct)}指标", level=1); cur = ct
        doc.add_heading(ind.name_cn or "", level=2)
        t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
        for label, key in FIELDS:
            if key == "_path":
                val = " / ".join(classification_path(db, ind.classification_id))
            elif key == "_tags":
                val = source_tags_str(ind)
            else:
                val = getattr(ind, key, "") or ""
            if not val:
                continue
            cells = t.add_row().cells
            cells[0].text = label; cells[1].text = str(val)
            cells[0].width = Cm(2.2); cells[1].width = Cm(14.8)
            for cell in cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10.5); run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=indicator_changes.docx"})

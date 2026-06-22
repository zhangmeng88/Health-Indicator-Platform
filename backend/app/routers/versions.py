"""最终成果版本管理：对当前指标池打快照，可留存、查看、删除并按版本导出。"""
import io

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from ..database import get_db
from ..security import require_admin
from ..models import OutputVersion, User
from ..schemas import VersionCreate, VersionOut
from ..utils import audit
from .export import _walk, _indicators, source_tags_str, HEADERS

router = APIRouter(prefix="/versions", tags=["版本管理"])


def _build_snapshot(db: Session) -> list[dict]:
    """按分类层级 + 同类内顺序，产出有序的指标快照行。"""
    rows = []
    for node, path in _walk(db):
        levels = (path + ["", "", ""])[:3]
        for ind in _indicators(db, node.id):
            rows.append({
                "source_title": ind.source_standard.title if ind.source_standard else "",
                "l1": levels[0], "l2": levels[1], "l3": levels[2],
                "identifier": ind.identifier or "", "name_cn": ind.name_cn or "",
                "name_en": ind.name_en or "", "unit": ind.unit or "",
                "definition": ind.definition or "", "method": ind.method or "",
                "description": ind.description or "", "survey_method": ind.survey_method or "",
                "data_source": ind.data_source or "", "frequency": ind.frequency or "",
                "stratification": ind.stratification or "",
                "source_tags_text": source_tags_str(ind),
                "indicator_type": ind.indicator_type or "",
            })
    return rows


def _out(v: OutputVersion) -> VersionOut:
    return VersionOut(id=v.id, label=v.label, note=v.note or "", indicator_count=v.indicator_count or 0,
                      creator_name=v.creator.display_name if v.creator else None, created_at=v.created_at)


@router.get("", response_model=list[VersionOut], summary="版本列表（管理员）")
def list_versions(db: Session = Depends(get_db), _: User = Depends(require_admin)):
    rows = db.query(OutputVersion).order_by(OutputVersion.id.desc()).all()
    return [_out(v) for v in rows]


@router.post("", response_model=VersionOut, status_code=201, summary="基于当前指标池创建新版本（管理员）")
def create_version(body: VersionCreate, db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    if not body.label.strip():
        raise HTTPException(400, "请填写版本名称")
    snap = _build_snapshot(db)
    v = OutputVersion(label=body.label.strip(), note=body.note or "", snapshot={"indicators": snap},
                      indicator_count=len(snap), created_by=admin.id)
    db.add(v); db.flush()
    audit(db, admin.id, "create_version", "version", v.id, {"label": v.label, "count": len(snap)})
    db.commit(); db.refresh(v)
    return _out(v)


@router.get("/{vid}", summary="版本详情（含快照，管理员）")
def get_version(vid: int, db: Session = Depends(get_db), _: User = Depends(require_admin)):
    v = db.get(OutputVersion, vid)
    if not v:
        raise HTTPException(404, "版本不存在")
    return {"id": v.id, "label": v.label, "note": v.note or "",
            "indicator_count": v.indicator_count or 0,
            "creator_name": v.creator.display_name if v.creator else None,
            "created_at": v.created_at, "indicators": (v.snapshot or {}).get("indicators", [])}


@router.delete("/{vid}", status_code=204, summary="删除版本（管理员）")
def delete_version(vid: int, db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    v = db.get(OutputVersion, vid)
    if not v:
        raise HTTPException(404, "版本不存在")
    audit(db, admin.id, "delete_version", "version", v.id, {"label": v.label})
    db.delete(v); db.commit()


def _rows_of(db, vid):
    v = db.get(OutputVersion, vid)
    if not v:
        raise HTTPException(404, "版本不存在")
    return v, (v.snapshot or {}).get("indicators", [])


@router.get("/{vid}/export/excel", summary="导出某版本 Excel（管理员）")
def export_version_excel(vid: int, db: Session = Depends(get_db), _: User = Depends(require_admin)):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    v, rows = _rows_of(db, vid)
    wb = Workbook(); ws = wb.active; ws.title = "卫生统计指标"
    thin = Side(style="thin", color="D0D0D0"); border = Border(thin, thin, thin, thin)
    ws.append(HEADERS)
    for c in range(1, len(HEADERS) + 1):
        cell = ws.cell(1, c)
        cell.font = Font(name="宋体", size=10, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E5F")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    for r in rows:
        ws.append([
            r.get("source_title", ""), r.get("l1", ""), r.get("l2", ""), r.get("l3", ""),
            r.get("identifier", ""), r.get("name_cn", ""), r.get("name_en", ""), r.get("unit", ""),
            r.get("definition", ""), r.get("method", ""), r.get("description", ""),
            r.get("survey_method", ""), r.get("data_source", ""), r.get("frequency", ""),
            r.get("stratification", ""), r.get("source_tags_text", ""), r.get("indicator_type", ""),
        ])
        rr = ws.max_row
        for c in range(1, len(HEADERS) + 1):
            cc = ws.cell(rr, c)
            cc.font = Font(name="宋体", size=10)
            cc.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cc.border = border
    for idx, w in enumerate([28, 12, 12, 12, 14, 22, 30, 8, 40, 32, 40, 10, 16, 10, 30, 24, 12], start=1):
        ws.column_dimensions[chr(64 + idx)].width = w
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": "attachment; filename=indicator_version.xlsx"})


@router.get("/{vid}/export/word", summary="导出某版本 Word（管理员）")
def export_version_word(vid: int, db: Session = Depends(get_db), _: User = Depends(require_admin)):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    v, rows = _rows_of(db, vid)
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2)
    doc.add_heading(f"卫生统计指标（{v.label}）", level=0)
    if v.note:
        doc.add_paragraph(v.note)

    FIELDS = [("标识符", "identifier"), ("英文名称", "name_en"), ("计量单位", "unit"),
              ("定义", "definition"), ("计算方法", "method"), ("指标说明", "description"),
              ("调查方法", "survey_method"), ("数据来源", "data_source"), ("发布频率", "frequency"),
              ("分层统计", "stratification"), ("来源标签", "source_tags_text"), ("指标类型", "indicator_type")]
    last_path = None
    for r in rows:
        path = " / ".join([p for p in [r.get("l1"), r.get("l2"), r.get("l3")] if p])
        if path != last_path:
            doc.add_heading(path or "未分类", level=1); last_path = path
        doc.add_heading(r.get("name_cn", ""), level=2)
        t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
        for label, key in FIELDS:
            val = r.get(key, "")
            if not val:
                continue
            cells = t.add_row().cells
            cells[0].text = label; cells[1].text = str(val)
            cells[0].width = Cm(2.2); cells[1].width = Cm(14.8)
            for cell in cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10.5)
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=indicator_version.docx"})
